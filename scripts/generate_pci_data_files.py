import json
import csv
import pandas as pd
from docx import Document
from fpdf import FPDF
from faker import Faker
import random

# Initialize Faker to generate random data
fake = Faker(['en_US'])

# Configuration
NUMBER_OF_DOCUMENTS = 10000  # Number of documents to generate
OUTPUT_FILE = "US_PCI_data1"  # Output file name without extension

# Function to create a fake document
def create_fake_document():
    return {
        "name": fake.name(),
        "email": fake.email(),
        "address": {
            "street": fake.street_address(),
            "city": fake.city(),
            "state": fake.state(),
            "zip": fake.zipcode()
        },
        "phone": fake.phone_number(),
        "credit_card_num": fake.credit_card_number(card_type=None),
        "credit_card_exp": fake.credit_card_expire(),
        "credit_card_security_num": fake.credit_card_security_code()
    }


# Generate and save the documents to a file
def save_documents_to_file(num_docs, filename, file_format='json'):
    documents = [create_fake_document() for _ in range(num_docs)]

    if file_format == 'json':
        with open(f"{filename}.json", 'w') as outfile:
            json.dump(documents, outfile, indent=4)
        print(f"Saved {num_docs} documents to {filename}.json")

    elif file_format == 'csv':
        with open(f"{filename}.csv", 'w', newline='') as outfile:
            fieldnames = documents[0].keys()
            writer = csv.DictWriter(outfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(documents)
        print(f"Saved {num_docs} documents to {filename}.csv")

    elif file_format == 'excel':
        df = pd.DataFrame(documents)
        df.to_excel(f"{filename}.xlsx", index=False)
        print(f"Saved {num_docs} documents to {filename}.xlsx")

    elif file_format == 'word':
        doc = Document()
        for document in documents:
            for key, value in document.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph("\n")
        doc.save(f"{filename}.docx")
        print(f"Saved {num_docs} documents to {filename}.docx")

    elif file_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for document in documents:
            for key, value in document.items():
                pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
            pdf.cell(200, 10, txt="\n", ln=True)
        pdf.output(f"{filename}.pdf")
        print(f"Saved {num_docs} documents to {filename}.pdf")

    elif file_format == 'txt':
        with open(f"{filename}.txt", 'w') as outfile:
            for document in documents:
                for key, value in document.items():
                    outfile.write(f"{key}: {value}\n")
                outfile.write("\n")
        print(f"Saved {num_docs} documents to {filename}.txt")

    else:
        raise ValueError("Invalid file format. Please use 'json', 'csv', 'excel', 'word', 'pdf', or 'txt'.")

# Generate and save the documents
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='json')  # Save as JSON
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='csv')  # Save as CSV
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='excel')  # Save as Excel
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='word')  # Save as Word
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='pdf')  # Save as PDF
save_documents_to_file(NUMBER_OF_DOCUMENTS, OUTPUT_FILE, file_format='txt')  # Save as Plain
