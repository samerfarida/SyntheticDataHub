# %%
import json
import csv
import pandas as pd
from docx import Document
from fpdf import FPDF
from faker import Faker
import random
import argparse
import string

# Initialize argument parser
parser = argparse.ArgumentParser(description="Generate IT/Security data")

# Add arguments with default values
parser.add_argument("--locale", default='en_US', help="Locale for Faker data generation (default: en_US)")
parser.add_argument("--records", type=int, default=10, help="Number of records to generate (default: 100)")

# Parse arguments
args = parser.parse_args()

# Use parsed arguments or default values
LOCAL = args.locale
NUMBER_OF_RECORDS = args.records

# Generate an 8-character random string
RANDOM_STRING = ''.join(random.choices(string.ascii_letters + string.digits, k=8))
# Generate output file name
OUTPUT_FILE = f"{LOCAL}_IT_security_data_{NUMBER_OF_RECORDS}_{RANDOM_STRING}"


# %%
# Initialize Faker to generate random data
fake = Faker([LOCAL])

# Function to create a fake IT/Security document
def create_fake_document():
    return {
        "user_id": fake.uuid4(),
        "username": fake.user_name(),
        "email": fake.email(),
        "department": fake.random_element(elements=('IT', 'Finance', 'HR', 'Marketing', 'Sales', 'Legal')),
        "job_title": fake.random_element(elements=('Security Analyst', 'Network Engineer', 'Software Developer', 'Data Scientist', 'IT Support')),
        "access_level": fake.random_element(elements=('Low', 'Medium', 'High')),
        "incident_type": fake.random_element(elements=('Unauthorized Access', 'Data Breach', 'Phishing Attack', 'Malware Infection', 'DDoS Attack')),
        "severity": fake.random_element(elements=('Low', 'Medium', 'High')),
        "incident_date": fake.date_time_this_year().strftime("%Y-%m-%d %H:%M:%S"),
        "incident_status": fake.random_element(elements=('Open', 'In Progress', 'Resolved', 'Closed')),
        "file_name": f"{fake.word()}.{random.choice(['docx', 'xlsx', 'pdf', 'txt', 'csv'])}",
        "file_size_kb": random.randint(10, 10000),
        "file_type": fake.random_element(elements=('Sensitive', 'Confidential', 'Public')),
        "file_action": fake.random_element(elements=('Accessed', 'Modified', 'Deleted', 'Shared', 'Downloaded')),
        "ip_address": fake.ipv4_private(),
        "mac_address": fake.mac_address(),
        "location": fake.city(),
    }



# %%
# Generate and save the documents to a file
def save_records_to_file(num_docs, filename, file_format='json'):
    documents = [create_fake_document() for _ in range(num_docs)]

    if file_format == 'json':
        with open(f"{filename}.json", 'w') as outfile:
            json.dump(documents, outfile, indent=4)
        print(f"Saved {num_docs} records to {filename}.json")

    elif file_format == 'csv':
        with open(f"{filename}.csv", 'w', newline='') as outfile:
            fieldnames = documents[0].keys()
            writer = csv.DictWriter(outfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(documents)
        print(f"Saved {num_docs} records to {filename}.csv")

    elif file_format == 'excel':
        df = pd.DataFrame(documents)
        df.to_excel(f"{filename}.xlsx", index=False)
        print(f"Saved {num_docs} records to {filename}.xlsx")

    elif file_format == 'word':
        doc = Document()
        for document in documents:
            for key, value in document.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph("\n")
        doc.save(f"{filename}.docx")
        print(f"Saved {num_docs} records to {filename}.docx")

    elif file_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for document in documents:
            for key, value in document.items():
                pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
            pdf.cell(200, 10, txt="\n", ln=True)
        pdf.output(f"{filename}.pdf")
        print(f"Saved {num_docs} records to {filename}.pdf")

    elif file_format == 'txt':
        with open(f"{filename}.txt", 'w') as outfile:
            for document in documents:
                for key, value in document.items():
                    outfile.write(f"{key}: {value}\n")
                outfile.write("\n")
        print(f"Saved {num_docs} records to {filename}.txt")

    else:
        raise ValueError("Invalid file format. Please use 'json', 'csv', 'excel', 'word', 'pdf', or 'txt'.")



# %%

# Generate and save the documents
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='json')  # Save as JSON
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='csv')  # Save as CSV
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='excel')  # Save as Excel
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='word')  # Save as Word
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='pdf')  # Save as PDF
save_records_to_file(NUMBER_OF_RECORDS, OUTPUT_FILE, file_format='txt')  # Save as Plain Text


